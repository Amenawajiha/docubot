import argparse
import gc
import logging
import sys
from pathlib import Path
from typing import List

from docx import Document

from src.ingestion.embedding_manager import EmbeddingManager
from src.ingestion.text_chunker import TextChunker
from src.utils.config_loader import get_config

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Orchestrates document loading, chunking, and storage in ChromaDB"""

    def __init__(self, file_path: str, batch_size: int = 1):
        """
        Initialize document processor

        Args:
            file_path: Path to document file
            batch_size: Batch size for embedding generation (default: 1, safest)
                       Use 1 for low memory, 2-4 only if you have 8GB+ RAM
        """
        self.chroma_db_path = get_config("vector.chroma_db_path")
        self.collection_name = get_config("vector.collection_name")
        self.embedding_model = get_config("vector.embedding_model_name")
        self.batch_size = batch_size

        # Initialize components
        self.chunker = TextChunker()
        self.embedding_manager = EmbeddingManager()

        # Convert file_path to Path object
        self.file_path = Path(file_path)

    def __read(self) -> str:
        """Load DOCX content"""
        try:
            doc = Document(self.file_path)

            # Extract all paragraphs
            content_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        content_parts.append(" | ".join(row_data))

            content = "\n\n".join(content_parts)

            if not content.strip():
                raise ValueError("No text content found in DOCX file")

            return content

        except Exception as e:
            raise RuntimeError(f"Error loading DOCX file: {str(e)}") from e

    def process_document_file(self) -> dict:
        """
        Process a single document file

        Returns:
            Dictionary with processing statistics
        """
        logger.info(f"\n{'='*60}")
        logger.info(f"Processing file: {self.file_path}")
        logger.info(f"{'='*60}\n")

        try:
            # 0. Delete existing chunks for this document (enables overwrite)
            existing_count = self.delete_document_chunks(self.file_path.name)
            if existing_count > 0:
                logger.info(f"Removed {existing_count} existing chunks for overwrite")

            # 1. Load document
            content = self.__read()
            logger.info(f"Loaded {len(content)} characters")

            # 2. Chunk document
            chunks = self.chunker.chunk(content)
            logger.info(f"Created {len(chunks)} chunks")

            # Free document content from memory immediately
            del content

            # 3. Embed chunks (chunks are dicts with 'text' key)
            chunk_texts = [chunk["text"] for chunk in chunks]
            embeddings = self.embedding_manager.embed_texts(
                chunk_texts, batch_size=self.batch_size
            )

            # Free chunk texts after embedding
            del chunk_texts

            # 4. Add to ChromaDB
            ids = []
            documents = []
            metadatas = []

            for chunk in chunks:
                doc_id = f"{self.file_path.stem}_{chunk['chunk_id']}"
                ids.append(doc_id)
                documents.append(chunk["text"])

                # Prepare metadata
                chunk_metadata = {
                    **chunk.get("metadata", {}),
                    "document_name": self.file_path.name,
                    "document_stem": self.file_path.stem,
                    "embedding_model": self.embedding_model,
                }
                metadatas.append(chunk_metadata)

            # Add to collection (embeddings already generated)
            self.collection.add(
                ids=ids,
                documents=documents,
                metadatas=metadatas,
                embeddings=(
                    embeddings.tolist() if hasattr(embeddings, "tolist") else embeddings
                ),
            )

            logger.info(f"Added {len(ids)} chunks to ChromaDB")

            return {
                "file": self.file_path.name,
                "status": "success",
                "chunks_created": len(chunks),
                "chunks_stored": len(ids),
                "content_size": sum(len(chunk["text"]) for chunk in chunks),
            }

        except Exception as e:
            logger.error(f"Error processing file {self.file_path}: {str(e)}")
            return {"file": self.file_path.name, "status": "error", "error": str(e)}

    def search(
        self, query: str, top_k: int = 5, include_metadata: bool = True
    ) -> List[dict]:
        """
        Search for relevant documents

        Args:
            query: Search query
            top_k: Number of results to return
            include_metadata: Whether to include metadata in results

        Returns:
            List of search results
        """
        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=top_k,
                include=["documents", "metadatas", "distances"],
            )

            formatted_results = []
            for i, (doc, metadata, distance) in enumerate(
                zip(
                    results["documents"][0],
                    results["metadatas"][0],
                    results["distances"][0],
                )
            ):
                formatted_results.append(
                    {
                        "rank": i + 1,
                        "content": doc,
                        "relevance_score": 1
                        - distance,  # Convert distance to similarity
                        "metadata": metadata if include_metadata else None,
                    }
                )

            return formatted_results

        except Exception as e:
            logger.error(f"Error searching: {str(e)}")
            raise

    def get_collection_stats(self) -> dict:
        """Get statistics about the collection"""
        try:
            count = self.collection.count()
            return {
                "collection_name": self.collection_name,
                "total_documents": count,
                "embedding_model": self.embedding_model,
                "chroma_db_path": self.chroma_db_path,
            }
        except Exception as e:
            logger.error(f"Error getting collection stats: {str(e)}")
            raise

    def delete_document_chunks(self, document_name: str) -> int:
        """
        Delete all chunks belonging to a specific document.

        Args:
            document_name: Name of the document (e.g., "Schengen Visa FAQs.docx")

        Returns:
            Number of chunks deleted
        """
        try:
            # Query to find all chunks with this document_name
            results = self.collection.get(
                where={"document_name": document_name}, include=["metadatas"]
            )

            ids_to_delete = results.get("ids", [])

            if ids_to_delete:
                self.collection.delete(ids=ids_to_delete)
                logger.info(
                    f"Deleted {len(ids_to_delete)} chunks for document: {document_name}"
                )
            else:
                logger.info(f"No existing chunks found for document: {document_name}")

            return len(ids_to_delete)

        except Exception as e:
            logger.error(f"Error deleting document chunks: {str(e)}")
            raise

    def list_documents(self) -> List[dict]:
        """
        List all unique documents in the collection.

        Returns:
            List of dicts with document_name and chunk_count
        """
        try:
            # Get all metadata
            results = self.collection.get(include=["metadatas"])
            metadatas = results.get("metadatas", [])

            # Count chunks per document
            doc_counts = {}
            for metadata in metadatas:
                doc_name = metadata.get("document_name", "unknown")
                doc_counts[doc_name] = doc_counts.get(doc_name, 0) + 1

            return [
                {"document_name": name, "chunk_count": count}
                for name, count in sorted(doc_counts.items())
            ]

        except Exception as e:
            logger.error(f"Error listing documents: {str(e)}")
            raise

    def get_document_chunk_count(self, document_name: str) -> int:
        """
        Get the number of chunks for a specific document.

        Args:
            document_name: Name of the document

        Returns:
            Number of chunks for this document
        """
        try:
            results = self.collection.get(
                where={"document_name": document_name},
                include=[],  # We only need IDs, not content
            )
            return len(results.get("ids", []))
        except Exception as e:
            logger.error(f"Error getting document chunk count: {str(e)}")
            return 0


def main():
    parser = argparse.ArgumentParser(description="Load documents into ChromaDB")

    parser.add_argument("--file", type=str, help="Path to single docx file")

    parser.add_argument(
        "--search", type=str, help="Search query after loading (optional)"
    )

    parser.add_argument(
        "--search-top-k",
        type=int,
        default=5,
        help="Number of results for search (default: 5)",
    )

    parser.add_argument(
        "--stats", action="store_true", help="Show collection statistics"
    )

    parser.add_argument(
        "--batch-size",
        type=int,
        default=1,
        help="Batch size for embedding (default: 1 = safest). Use 2-4 only with 8GB+ RAM.",
    )

    parser.add_argument(
        "--list-docs", action="store_true", help="List all documents in the collection"
    )

    parser.add_argument(
        "--delete-doc", type=str, help="Delete all chunks for a specific document"
    )

    args = parser.parse_args()

    # Validate arguments
    if (
        not args.file
        and not args.search
        and not args.stats
        and not args.list_docs
        and not args.delete_doc
    ):
        parser.print_help()
        sys.exit(1)

    try:
        # Initialize processor (use dummy path for non-file operations)
        file_path = args.file if args.file else "dummy.docx"
        processor = DocumentProcessor(file_path=file_path, batch_size=args.batch_size)

        # List documents
        if args.list_docs:
            docs = processor.list_documents()
            print(f"\n📚 Documents in collection:")
            if docs:
                for doc in docs:
                    print(f"   • {doc['document_name']} ({doc['chunk_count']} chunks)")
            else:
                print("   No documents found")

        # Delete document
        if args.delete_doc:
            deleted = processor.delete_document_chunks(args.delete_doc)
            print(f"\n🗑️  Deleted {deleted} chunks for: {args.delete_doc}")

        # Process files
        if args.file:
            logger.info(f"Loading single file: {args.file}")
            result = processor.process_document_file()
            print(f"\n✅ File processed: {result}")

        # Show statistics
        if args.stats:
            stats = processor.get_collection_stats()
            print(f"\n📊 Collection Statistics:")
            print(f"   Collection: {stats['collection_name']}")
            print(f"   Total documents: {stats['total_documents']}")
            print(f"   Embedding model: {stats['embedding_model']}")
            print(f"   Storage path: {stats['chroma_db_path']}")

        # Perform search
        if args.search:
            logger.info(f"Searching for: {args.search}")
            results = processor.search(args.search, top_k=args.search_top_k)

            print(f"\n🔍 Search Results for: '{args.search}'")
            for result in results:
                print(
                    f"\n   Rank {result['rank']} (Score: {result['relevance_score']:.4f})"
                )
                print(f"   Content: {result['content'][:200]}...")
                if result["metadata"]:
                    print(
                        f"   Source: {result['metadata'].get('document_name', 'Unknown')}"
                    )

    except Exception as e:
        logger.error(f"Error: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()

"""
### Load Single DOCX File
```bash
python document_processor.py --file Schengen-Visa-FAQs.docx --chroma-path ./chroma_data
```

### Load Directory
```bash
python document_processor.py --dir ./documents/ --chroma-path ./chroma_data --chunk-size 1000
```

### Search Documents
```bash
python document_processor.py \
  --search "What documents do I need for Schengen visa?" \
  --chroma-path ./chroma_data \
  --search-top-k 5
```

### Show Statistics
```bash
python document_processor.py --stats --chroma-path ./chroma_data
```

### Combined: Load + Search
```bash
python document_processor.py \
  --file Schengen-Visa-FAQs.docx \
  --search "flight itinerary validity" \
  --chroma-path ./chroma_data \
  --stats
```
## 11. Programmatic Usage

```python
# Use in your services


# Initialize
processor = DocumentProcessor(
    chroma_db_path="./chroma_data",
    collection_name="visa_documents",
    chunk_size=800,
    chunk_overlap=200
)

# Load documents
result = processor.process_file("Schengen-Visa-FAQs.docx")
print(f"Loaded {result['chunks_created']} chunks")

# Search
results = processor.search("What is flight itinerary?", top_k=3)
for result in results:
    print(f"Score: {result['relevance_score']:.4f}")
    print(f"Content: {result['content']}")
    print(f"Source: {result['metadata']['document_name']}")

"""
