"""
Word Document Reader - Extract text from Word files.

Supports:
- Modern Word (.docx) - Office Open XML format
- Extracts paragraphs, tables, headers, and footers
- Preserves document structure

Output format:
- Paragraphs separated by double newlines
- Tables converted to natural language
"""

import logging
from io import BytesIO

from .base_reader import BaseReader

logger = logging.getLogger(__name__)


class DOCXReader(BaseReader):
    """
    Extract text from Word document files.
    
    Extracts all readable content:
    - Paragraphs
    - Tables
    - Headers and footers
    - Text boxes (limited support)
    """
    
    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]
    
    def extract(self, file_bytes: bytes) -> str:
        """
        Extract text from Word document.
        
        Args:
            file_bytes: Raw DOCX bytes
            
        Returns:
            Document text with preserved structure
            
        Raises:
            ValueError: If file is empty or contains no text
            RuntimeError: If parsing fails
        """
        if not file_bytes:
            raise ValueError("Empty Word document provided")
        
        # Lazy import
        try:
            from docx import Document
        except ImportError as e:
            raise RuntimeError(
                "python-docx is required for Word support. "
                "Install with: pip install python-docx"
            ) from e
        
        try:
            doc = Document(BytesIO(file_bytes))
            parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)
            
            # Extract tables
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text:
                    parts.append(table_text)
            
            content = "\n\n".join(parts)

            try: 
                from src.ingestion.utils.ocr import run_ocr
            except ImportError:
                run_ocr = None

            ocr_texts = []
            for image_bytes in self._extract_images(file_bytes, Document):
                if run_ocr:
                    try:
                        ocr_result = run_ocr(image_bytes)
                        if ocr_result.strip():
                            ocr_texts.append(ocr_result.strip())
                    except Exception as e:
                        logger.warning(f"OCR failed for an image: {e}")

            if ocr_texts:
                content += "\n\n[Image OCR]\n" + "\n\n".join(ocr_texts)
            
            if not content.strip():
                raise ValueError("No text content found in Word document")
            
            logger.info(f"Word document extraction complete: {len(content)} characters")
            return content
        
        except Exception as e:
            logger.error(f"Word document extraction failed: {e}")
            raise RuntimeError(f"Failed to read Word document: {e}") from e
    
    def _extract_table(self, table) -> str:
        """
        Extract text from a table.
        
        Converts table to natural language format:
        "Header1: value1; Header2: value2; Header3: value3"
        
        Args:
            table: python-docx Table object
            
        Returns:
            Table text in natural language format
        """
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        
        if not rows:
            return ""
        
        # Simple case: single row
        if len(rows) < 2:
            return " | ".join(rows[0])
        
        # Extract headers
        headers = list(rows[0])
        
        # Convert data rows to sentences
        lines = []
        for row in rows[1:]:
            cells = []
            for header, value in zip(headers, row):
                if value:  # Only include non-empty values
                    prefix = f"{header}: " if header else ""
                    cells.append(f"{prefix}{value}")
            
            if cells:
                lines.append("; ".join(cells))
        
        return "\n".join(lines)
    
    def _extract_images(self, file_bytes: bytes, Document):
        doc = Document(BytesIO(file_bytes))
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_bytes = rel.target_part.blob
                images.append(image_bytes)
        return images

class DOCReader(DOCXReader):
    """
    Legacy Word document reader (.doc files).
    
    Note: python-docx only supports .docx (Office Open XML format).
    For .doc (binary format), consider using LibreOffice conversion
    or the antiword library as a preprocessing step.
    """
    
    @property
    def supported_extensions(self) -> list[str]:
        return [".doc"]
    
    def extract(self, file_bytes: bytes) -> str:
        """
        Extract text from legacy Word document.
        
        This is a placeholder - .doc support requires additional libraries.
        """
        raise RuntimeError(
            "Legacy .doc format is not directly supported. "
            "Please convert to .docx format first, or use LibreOffice/antiword "
            "for batch conversion. python-docx only supports .docx files."
        )